"""
parser.py  –  Resume field extractor
========================================
Extracts: name, email, mobile, qualification, passing year,
          latest/current company, total experience.

Passing year  : takes end year of the range on the same line first,
                then searches the next 2 lines.
Latest company: four-priority cascade (current flag → latest start_year
                → first record → full-text fallback).
Total exp.    : parses ALL date ranges, merges overlapping intervals,
                sums and formats as "X years Y months".
"""

import re
import datetime
from typing import Optional

import pdfplumber
import spacy
from docx import Document
import os

# Load once at module level
nlp = spacy.load("en_core_web_sm")

# ══════════════════════════════════════════════════════════════════
# SHARED MONTH / DATE CONSTANTS
# ══════════════════════════════════════════════════════════════════

_MONTH_MAP = {
    'jan': 1, 'january': 1,
    'feb': 2, 'february': 2,
    'mar': 3, 'march': 3,
    'apr': 4, 'april': 4,
    'may': 5,
    'jun': 6, 'june': 6,
    'jul': 7, 'july': 7,
    'aug': 8, 'august': 8,
    'sep': 9, 'sept': 9, 'september': 9,
    'oct': 10, 'october': 10,
    'nov': 11, 'november': 11,
    'dec': 12, 'december': 12,
}

_MONTHS_PAT = (
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|'
    r'Dec(?:ember)?)'
)

_CURRENT_PAT = (
    r'(?:Present|Current|Till[\s]*Date|Till[\s]*Now|Ongoing|'
    r'Now|Running|Continuing|Cont\.?)'
)

# ── Pattern A: "Jan 2020 – Present" / "Jan'20 – Dec'22" ──────────
_PAT_A = re.compile(
    r'(' + _MONTHS_PAT + r")" + r"['\s.,\-]{0,3}(\d{2,4})"
    r'\s*(?:–|—|-|to|till)\s*'
    r'(?:(' + _MONTHS_PAT + r")['\s.,\-]{0,3}(\d{2,4})"
    r'|(' + _CURRENT_PAT + r'))',
    re.IGNORECASE
)

# ── Pattern B: "2020 – Present" / "2019 – 2022" ──────────────────
_PAT_B = re.compile(
    r'\b(20\d{2}|19\d{2})\s*(?:–|—|-|to|till)\s*'
    r'(?:(20\d{2}|19\d{2})|(' + _CURRENT_PAT + r'))\b',
    re.IGNORECASE
)

# ── Pattern C: "03/2018 – 12/2022" (MM/YYYY) ─────────────────────
_PAT_C = re.compile(
    r'(0[1-9]|1[0-2])/(20\d{2}|19\d{2})'
    r'\s*(?:–|—|-|to)\s*'
    r'(?:(0[1-9]|1[0-2])/(20\d{2}|19\d{2})'
    r'|(' + _CURRENT_PAT + r'))',
    re.IGNORECASE
)

_CURRENT_RE = re.compile(r'\b' + _CURRENT_PAT + r'\b', re.IGNORECASE)

# Valid year range for resumes
_YEAR_RANGE_RE  = re.compile(r'\b(19[89]\d|20[0-2]\d)\s*[-–—]+\s*(19[89]\d|20[0-2]\d)\b')
_YEAR_SINGLE_RE = re.compile(r'\b(19[89]\d|20[0-2]\d)\b')


# ══════════════════════════════════════════════════════════════════
# PDF TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════

def extract_text(file_path: str) -> str:

    ext = os.path.splitext(file_path)[1].lower()

    # PDF
    if ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    # DOCX
    elif ext == ".docx":

        doc = Document(file_path)

        text = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_text))

        return "\n".join(text)

    return ""


# ══════════════════════════════════════════════════════════════════
# BASIC FIELDS
# ══════════════════════════════════════════════════════════════════

def extract_name(text):

    lines = text.split("\n")

    for line in lines[:10]:

        line = line.strip()

        if 3 < len(line) < 50:

            if not re.search(r'@|\d', line):

                return line

    return "Not Found"



def extract_email(text):

    pattern = r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}'

    match = re.search(pattern, text)

    if match:
        return match.group()

    return "Not Found"


def extract_mobile(text):

    pattern = r'(\+91[\s-]?)?[6-9]\d{9}'

    match = re.search(pattern, text)

    if match:
        return match.group()

    return "Not Found"


# ══════════════════════════════════════════════════════════════════
# SECTION SLICER
# ══════════════════════════════════════════════════════════════════

def _slice_section(text: str, start_kws: list, stop_kws: list) -> str:
    lines = text.split("\n")
    start = None
    end   = len(lines)
    for i, line in enumerate(lines):
        up = line.upper().strip()
        if start is None:
            # Require keyword to be the *dominant* content of the line
            if any(kw in up for kw in start_kws) and len(up) < 60:
                start = i
        else:
            if any(kw in up for kw in stop_kws) and len(up) < 60:
                end = i
                break
    return "" if start is None else "\n".join(lines[start:end])


# ── Section header lists ──────────────────────────────────────────

_EDU_START = [
    "EDUCATION", "ACADEMIC QUALIFICATION", "ACADEMIC QUALIFICATIONS",
    "QUALIFICATION", "EDUCATIONAL QUALIFICATION", "ACADEMIC BACKGROUND",
    "ACADEMIC DETAILS", "EDUCATIONAL DETAILS",
]
_EDU_STOP = [
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
    "EMPLOYMENT", "CAREER HISTORY", "SKILLS", "TECHNICAL SKILLS",
    "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS", "AWARDS",
    "INTERNSHIP", "INTERNSHIPS", "DECLARATION",
]

_EXP_START = [
    "WORK EXPERIENCE", "EXPERIENCE", "PROFESSIONAL EXPERIENCE",
    "EMPLOYMENT HISTORY", "CAREER HISTORY", "INTERNSHIP", "INTERNSHIPS",
    "WORK HISTORY",
]
_EXP_STOP = [
    "EDUCATION", "ACADEMIC", "SKILLS", "TECHNICAL SKILLS",
    "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS", "AWARDS",
    "DECLARATION", "HOBBIES", "PERSONAL DETAILS", "LANGUAGES",
    "REFERENCES", "OBJECTIVE", "SUMMARY",
]


def extract_education_section(text: str) -> str:
    return _slice_section(text, _EDU_START, _EDU_STOP)


def extract_experience_section(text: str) -> str:
    return _slice_section(text, _EXP_START, _EXP_STOP)


# ══════════════════════════════════════════════════════════════════
# QUALIFICATION + PASSING YEAR  (most accurate logic)
# ══════════════════════════════════════════════════════════════════

_QUALIFICATION_RANK = [
    # (display_name, rank, regex_pattern)
    ("PhD",                        9, r'\bph\.?\s*d\.?\b'),
    ("Doctorate",                  9, r'\bdoctorate\b'),
    ("MBA",                        8, r'\bm\.?\s*b\.?\s*a\.?\b'),
    ("M.Tech",                     7, r'\bm\.?\s*tech\b'),
    ("M.E",                        7, r'\bm\.?\s*e\.?\b(?!ng|l|c|d|ch)'),
    ("M.Sc",                       7, r'\bm\.?\s*sc(?:ience)?\b'),
    ("MCA",                        7, r'\bmca\b'),
    ("Master of Science",          7, r'\bmaster\s+of\s+science\b'),
    ("Master of Computer Science", 7, r'\bmaster\s+of\s+computer\s+science\b'),
    ("B.Tech",                     6, r'\bb\.?\s*tech\b'),
    ("Bachelor of Technology",     6, r'\bbachelor\s+of\s+technology\b'),
    ("B.E",                        6, r'\bb\.?\s*e\.?\b(?!ng|l|c|d|ch|low|fore)'),
    ("Bachelor of Engineering",    6, r'\bbachelor\s+of\s+engineering\b'),
    ("BCA",                        5, r'\bbca\b'),
    ("B.Sc",                       5, r'\bb\.?\s*sc(?:ience)?\b'),
    ("B.Com",                      5, r'\bb\.?\s*com\b'),
    ("Diploma",                    4, r'\bdiploma\b'),
    ("12th / HSC",                 3, r'\b(?:12th|hsc|higher\s+secondary)\b'),
    ("10th / SSC",                 2, r'\b(?:10th|ssc|matriculation|secondary)\b'),
]


def _extract_passing_year_from_window(lines: list, idx: int) -> Optional[int]:
    """
    Best-effort year extraction around line idx.

    Priority:
      1. Year RANGE on same line  → take end year (completion year)
      2. Single year on same line
      3. Year range on lines idx+1 / idx+2 → take end year
      4. Single year on lines idx+1 / idx+2
    """
    same_line = lines[idx]
    # Priority 1 – range on same line
    m = _YEAR_RANGE_RE.search(same_line)
    if m:
        return int(m.group(2))   # end year

    # Priority 2 – single year on same line
    singles = _YEAR_SINGLE_RE.findall(same_line)
    if singles:
        return max(int(y) for y in singles)

    # Priority 3 & 4 – look at next two lines
    for j in range(idx + 1, min(len(lines), idx + 3)):
        next_line = lines[j]
        m2 = _YEAR_RANGE_RE.search(next_line)
        if m2:
            return int(m2.group(2))
        singles2 = _YEAR_SINGLE_RE.findall(next_line)
        if singles2:
            return max(int(y) for y in singles2)

    return None


def extract_highest_qualification_details(text: str) -> dict:
    lines = text.split("\n")
    found = []

    for display, rank, pat in _QUALIFICATION_RANK:
        for i, line in enumerate(lines):
            if re.search(pat, line, re.IGNORECASE):
                year = _extract_passing_year_from_window(lines, i)
                found.append({"qualification": display, "rank": rank, "passing_year": year})
                break   # one match per level is enough

    if not found:
        return {"qualification": "Not Found", "passing_year": "Not Found"}

    best = max(found, key=lambda x: x["rank"])
    return {
        "qualification": best["qualification"],
        "passing_year":  str(best["passing_year"]) if best["passing_year"] else "Not Found",
    }


def extract_qualification(text: str) -> str:
    return extract_highest_qualification_details(text)["qualification"]


def extract_passing_year(text: str) -> str:
    return extract_highest_qualification_details(text)["passing_year"]


# ══════════════════════════════════════════════════════════════════
# COMPANY DETECTION ENGINE
# ══════════════════════════════════════════════════════════════════

_COMPANY_SUFFIXES = [
    "pvt ltd", "pvt. ltd", "pvt.ltd", "private limited", "private ltd",
    "limited", "ltd.", " ltd", "llp", "inc.", " inc", "corp.", " corp",
    "corporation", " technologies", " solutions", " systems", " services",
    " group", " consulting", " infotech", " software", " techno",
]

_KNOWN_COMPANIES = {c.lower(): c for c in [
    # IT / Software
    "TCS", "Tata Consultancy Services", "Infosys", "Wipro", "HCL",
    "HCL Technologies", "Tech Mahindra", "Cognizant", "Capgemini",
    "Accenture", "IBM", "Oracle", "SAP", "Microsoft", "Google", "Amazon",
    "LTIMindtree", "Mphasis", "Persistent Systems", "Hexaware",
    "NIIT Technologies", "Mindtree", "Zensar", "KPIT", "Cyient",
    "Birlasoft", "L&T Infotech", "Mastech", "Patni", "iGate",
    # Engineering / Manufacturing
    "Larsen & Toubro", "L&T", "Tata Motors", "Mahindra", "Mahindra & Mahindra",
    "Ashok Leyland", "Bharat Forge", "Cummins", "Bosch", "Siemens", "ABB",
    "Schneider Electric", "GE", "Honeywell", "Eaton", "Emerson", "Parker",
    "Atlas Copco", "SKF", "Thermax", "ISRO", "BARC",
    # Defence / PSU
    "HAL", "BEL", "BHEL", "BEML", "DRDO", "NTPC", "ONGC", "SAIL",
    "Tata Advanced Systems", "Adani Defence", "Solar Industries",
    "Bharat Dynamics", "Ordnance Factory",
    # Banking / Finance
    "HDFC Bank", "ICICI Bank", "Axis Bank", "Kotak Mahindra Bank",
    "State Bank of India", "SBI", "Standard Chartered", "HSBC",
    "Bajaj Finserv", "IDFC", "Yes Bank", "IndusInd Bank",
    # Consulting
    "Deloitte", "EY", "PwC", "KPMG", "McKinsey", "BCG", "Bain",
    "A.T. Kearney", "Roland Berger",
    # Product / Startup
    "Zoho", "Freshworks", "Paytm", "PhonePe", "Razorpay",
    "Flipkart", "Swiggy", "Zomato", "Ola", "Nykaa", "BYJU",
]}

_INVALID_TOKENS = {
    # job titles
    "engineer", "manager", "developer", "analyst", "intern", "executive",
    "consultant", "technician", "officer", "associate", "specialist",
    "director", "lead", "head", "senior", "junior",
    # skill names
    "python", "sql", "aws", "azure", "html", "css", "javascript",
    "react", "linux", "tcp", "dns", "dhcp", "networking", "database",
    "machine learning", "deep learning", "data science",
    # section names
    "skills", "projects", "education", "certifications", "training",
    "responsibilities", "achievements",
    # institutions
    "university", "college", "institute", "school", "iit", "nit",
    # misc false positives
    "responsible", "experience", "present", "current", "profile",
}


def _is_valid_company(name: str) -> bool:
    if not name:
        return False
    name = name.strip()
    if len(name) < 3 or len(name.split()) > 10:
        return False
    # strip trailing punctuation noise
    name_clean = re.sub(r'[|/\\:,]+$', '', name).strip()
    low = name_clean.lower()
    if any(tok in low for tok in _INVALID_TOKENS):
        return False
    # Must have at least one alphabetic character
    if not re.search(r'[A-Za-z]', name_clean):
        return False
    return True


def _find_company_in_lines(lines: list) -> Optional[str]:
    """
    Search ordered list of text lines for a company name.
    Layer 1 → suffix match, Layer 2 → known dict, Layer 3 → spaCy ORG.
    """
    # Layer 1 – company suffix heuristic
    for line in lines:
        low = line.lower()
        if any(suf in low for suf in _COMPANY_SUFFIXES):
            # Try to return just the company part (strip trailing date/title)
            candidate = re.split(r'\s{2,}|\|', line)[0].strip()
            if _is_valid_company(candidate):
                return candidate

    # Layer 2 – known company dictionary
    combined = " ".join(lines)
    low_combined = combined.lower()
    for key, canonical in _KNOWN_COMPANIES.items():
        if key in low_combined:
            return canonical

    # Layer 3 – spaCy ORG NER
    doc = nlp(combined[:400])
    for ent in doc.ents:
        if ent.label_ == "ORG" and _is_valid_company(ent.text):
            return ent.text.strip()

    return None


# ══════════════════════════════════════════════════════════════════
# DATE RANGE PARSER  (returns structured intervals)
# ══════════════════════════════════════════════════════════════════

def _normalize_year(y: str) -> int:
    """Convert 2-digit or 4-digit year string to int."""
    y = int(y)
    return y + 2000 if y < 100 else y


def _parse_date_range(line: str) -> Optional[dict]:
    """
    Try all three patterns against a single line.
    Returns dict with start_year, start_month, end_year, end_month, is_current
    or None if no date range found.
    """
    now = datetime.datetime.now()

    # ── Pattern A: "Jan 2020 – Present" ─────────────────────────
    m = _PAT_A.search(line)
    if m:
        sm = _MONTH_MAP.get(m.group(1).lower()[:3], 1)
        sy = _normalize_year(m.group(2))
        current = bool(m.group(5))
        if current:
            em, ey = now.month, now.year
        else:
            em = _MONTH_MAP.get((m.group(3) or "jan").lower()[:3], 1)
            ey = _normalize_year(m.group(4)) if m.group(4) else now.year
        return dict(start_year=sy, start_month=sm,
                    end_year=ey, end_month=em, is_current=current)

    # ── Pattern B: "2020 – Present" / "2019 – 2022" ─────────────
    m = _PAT_B.search(line)
    if m:
        sy = int(m.group(1))
        current = bool(m.group(3))
        ey = now.year if current else int(m.group(2))
        return dict(start_year=sy, start_month=1,
                    end_year=ey, end_month=now.month if current else 12,
                    is_current=current)

    # ── Pattern C: "03/2018 – 12/2022" ──────────────────────────
    m = _PAT_C.search(line)
    if m:
        sm, sy = int(m.group(1)), int(m.group(2))
        current = bool(m.group(5))
        if current:
            em, ey = now.month, now.year
        else:
            em = int(m.group(3)) if m.group(3) else 12
            ey = int(m.group(4)) if m.group(4) else now.year
        return dict(start_year=sy, start_month=sm,
                    end_year=ey, end_month=em, is_current=current)

    return None


# ══════════════════════════════════════════════════════════════════
# EXPERIENCE RECORDS  (company + dates per job)
# ══════════════════════════════════════════════════════════════════

def extract_experience_records(experience_text: str) -> list:
    """
    Returns list of dicts:
      company, start_year, start_month, end_year, end_month, is_current
    """
    records = []
    lines   = [l.rstrip() for l in experience_text.split("\n")]
    n       = len(lines)

    for i, line in enumerate(lines):
        dr = _parse_date_range(line)
        if dr is None:
            continue

        # ── Build a search window: look ABOVE first (most common layout),
        #    then same line prefix, then below ────────────────────────────
        above = [lines[j] for j in range(max(0, i - 3), i) if lines[j].strip()]
        above.reverse()   # closest line first

        # Same line text BEFORE the date (company on same line)
        prefix = re.split(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|'
                          r'19\d{2}|20\d{2})', line, maxsplit=1, flags=re.IGNORECASE)[0].strip()
        same_line_candidate = [prefix] if prefix else []

        below = [lines[j] for j in range(i + 1, min(n, i + 4)) if lines[j].strip()]

        # Priority: above (1-3 lines) → same-line prefix → below
        search_order = above + same_line_candidate + below

        company = _find_company_in_lines(search_order)

        if company and not any(r["company"] == company for r in records):
            records.append({
                "company":     company,
                "start_year":  dr["start_year"],
                "start_month": dr["start_month"],
                "end_year":    dr["end_year"],
                "end_month":   dr["end_month"],
                "is_current":  dr["is_current"],
                # keep old key for backward compat
                "current":     dr["is_current"],
            })

    return records


def extract_companies_from_experience(experience_text: str) -> list:
    """De-duped list of ORG names found by spaCy (for debug/audit)."""
    doc  = nlp(experience_text)
    seen = []
    for ent in doc.ents:
        if ent.label_ == "ORG":
            name = ent.text.strip()
            if len(name) > 2 and name not in seen:
                seen.append(name)
    return seen


# ══════════════════════════════════════════════════════════════════
# LATEST / CURRENT COMPANY
# ══════════════════════════════════════════════════════════════════

def extract_latest_company(text: str) -> str:
    """
    Priority cascade:
      1. Record explicitly flagged is_current (Present / Till Date …)
      2. Record with highest start_year
      3. First record in the list
      4. Full-text _find_company_in_lines fallback
    """
    exp_text = extract_experience_section(text)

    if not exp_text:
        result = _find_company_in_lines(text.split("\n"))
        return result if result else "Not Found"

    records = extract_experience_records(exp_text)

    # 1. Explicit current
    for r in records:
        if r["is_current"]:
            return r["company"]

    # 2. Most recent by start_year
    dated = [r for r in records if r.get("start_year")]
    if dated:
        return max(dated, key=lambda r: (r["start_year"], r["start_month"]))["company"]

    # 3. First record
    if records:
        return records[0]["company"]

    # 4. Full-text fallback
    result = _find_company_in_lines(exp_text.split("\n"))
    return result if result else "Not Found"


# ══════════════════════════════════════════════════════════════════
# TOTAL EXPERIENCE  (brand-new feature)
# ══════════════════════════════════════════════════════════════════

def _abs_month(year: int, month: int) -> int:
    """Convert (year, month) to an absolute month count."""
    return year * 12 + month


def _merge_intervals(intervals: list) -> list:
    """Merge a list of (start, end) absolute-month tuples."""
    if not intervals:
        return []
    intervals = sorted(intervals)
    merged = [list(intervals[0])]
    for s, e in intervals[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return merged


def _format_experience(total_months: int) -> str:
    if total_months <= 0:
        return "Less than 1 month"
    years  = total_months // 12
    months = total_months % 12
    parts  = []
    if years:
        parts.append(f"{years} year{'s' if years > 1 else ''}")
    if months:
        parts.append(f"{months} month{'s' if months > 1 else ''}")
    return " ".join(parts)


def extract_total_experience(text: str) -> str:
    """
    Parse every date range inside the experience section,
    merge overlapping intervals, sum total months, and
    return a human-readable string like "3 years 4 months".
    """
    exp_text = extract_experience_section(text)
    if not exp_text:
        return "Not Found"

    lines     = exp_text.split("\n")
    intervals = []

    for line in lines:
        dr = _parse_date_range(line)
        if dr is None:
            continue
        s = _abs_month(dr["start_year"], dr["start_month"])
        e = _abs_month(dr["end_year"],   dr["end_month"])
        if e > s:
            intervals.append((s, e))

    if not intervals:
        return "Not Found"

    merged      = _merge_intervals(intervals)
    total_months = sum(e - s for s, e in merged)
    return _format_experience(total_months)

